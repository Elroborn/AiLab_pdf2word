import suanpan
from suanpan.app import app
from suanpan.app.arguments import Folder,File
import os
from io import StringIO
from io import open
from concurrent.futures import ProcessPoolExecutor

from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import process_pdf
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from docx import Document
from glob import glob

def read_from_pdf(file_path):
    with open(file_path, 'rb') as file:
        resource_manager = PDFResourceManager()
        return_str = StringIO()
        lap_params = LAParams()

        device = TextConverter(
            resource_manager, return_str, laparams=lap_params)
        process_pdf(resource_manager, device, file)
        device.close()

        content = return_str.getvalue()
        return_str.close()
        return content


def save_text_to_word(content, file_path):
    doc = Document()
    for line in content.split('\n'):
        paragraph = doc.add_paragraph()
        paragraph.add_run(remove_control_characters(line))
    doc.save(file_path)


def remove_control_characters(content):
    mpa = dict.fromkeys(range(32))
    return content.translate(mpa)


def pdf_to_word(pdf_file_path, word_file_path):
    content = read_from_pdf(pdf_file_path)
    save_text_to_word(content, word_file_path)



@app.input(Folder(key="inputData1"))
@app.output(Folder(key="outputData1"))
def pdf2word(context):
    args = context.args
    print("*"*20)
    print(args.inputData1)
    file_list = glob("%s/*.pdf"%args.inputData1) + glob("%s/*/*.pdf"%args.inputData1)
    print(file_list)


    tasks = []
    with ProcessPoolExecutor(max_workers=5) as executor:
        for file in file_list:
            extension_name = os.path.splitext(file)[-1]
            if extension_name != '.pdf':
                continue
            file_name = os.path.splitext(file)[0].split('/')[-1]

            word_file = "./res" + '/' + file_name + '.docx'
            print('processing', file)
            result = executor.submit(pdf_to_word, file, word_file)
            tasks.append(result)
    while True:
        exit_flag = True
        for task in tasks:
            if not task.done():
                exit_flag = False
        if exit_flag:
            print('done')
            return "res"




if __name__ == "__main__":
    suanpan.run(app)
